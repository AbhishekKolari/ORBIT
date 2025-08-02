import os
import logging
from docx import Document
from docx.shared import Inches
from PIL import Image
from transformers import Gemma3ForConditionalGeneration, AutoProcessor
import torch
from datetime import datetime
import re

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('gemma3_processing.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Image themes list
image_themes = ['kitchen', 'zoo', 'biking', 'tools', 'meeting room', 'reptile zoo', 'market', 'tech', 'gym', 
                'bedroom', 'classroom', 'garage', 'beach cleanup', 'camping', 'gardening', 'library', 
                'wardrobe', 'salon', 'construction', 'driveway', 'laboratory', 'home setup', 'urban', 
                'park', 'picnic', 'farm', 'bustop']

def get_theme_classification_messages(image_path):
    """Get messages for theme classification"""
    theme_prompt = f"""Look at this image and classify it into one of these themes: {', '.join(image_themes)}

Select the single most appropriate theme from the list that best describes the main subject or setting of the image. 
Respond with ONLY the theme name, nothing else."""
    
    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": "You are a helpful assistant that classifies images into predefined themes."}]
        },
        {
            "role": "user",
            "content": [
                {"type": "image", "image": image_path},
                {"type": "text", "text": theme_prompt}
            ]
        }
    ]
    return messages

def get_messages(idx, img_file, image_path):
    prompt_text = """Property dimension definitions:
Physical Properties: Materials (wood, metal, glass, plastic), States (solid, liquid, fragile, flexible), Structural characteristics (has wheels, has handle, has legs)
Taxonomic Properties: Biological categories (mammals, reptiles, birds), Artifact categories (furniture, tools, vehicles), Food categories (fruits, vegetables, grains)
Functional Properties: Use cases (can be worn, can hold liquid, can cut), Affordances (graspable, openable, foldable), Energy requirements (needs electricity, manual power, battery-operated)
Relational Properties: Spatial relations (items on top of other items, inside containers), grouping relations (couple, flock)

Generate exactly three creative counting questions for this image:
1. One question about PHYSICAL or TAXONOMIC properties
2. One question about FUNCTIONAL or RELATIONAL properties  
3. One counterfactual question about PHYSICAL or TAXONOMIC or FUNCTIONAL or RELATIONAL properties

Format each question EXACTLY like this:
1. [Question]? [Number] ([property]) ([objects])
2. [Question]? [Number] ([property]) ([objects])  
3. [Question]? [Number] ([property]) ([objects])

Example format:
1. How many objects made of wood are present? 2 (physical) (bowls)
2. How many breakable objects are visible? 1 (functional) (ceramic mug)
3. If the picture frames are removed from the wall, how many objects are visible hanging from the wall? 3 (relational) (kitchen knives)

Follow this exact format with no additional text or explanations."""
    
    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": "You are a helpful assistant that analyzes images and generates counting questions."}]
        },
        {
            "role": "user",
            "content": [
                {"type": "image", "image": image_path},
                {"type": "text", "text": prompt_text}
            ]
        }
    ]
    return messages

def validate_image(image_path):
    """Validate if the image can be opened and processed"""
    try:
        with Image.open(image_path) as img:
            img.verify()  # Verify it's a valid image
        # Re-open after verify (verify() closes the file)
        with Image.open(image_path) as img:
            img.convert("RGB")  # Test conversion
        return True
    except Exception as e:
        logger.error(f"Invalid image {image_path}: {e}")
        return False

def run_vlm_on_image(image_path, messages, model, processor, device):
    try:
        # Load the image
        image = Image.open(image_path).convert("RGB")
        
        # Update the messages with the actual image object
        messages[1]["content"][0]["image"] = image
        
        logger.info(f"Processing image with VLM: {os.path.basename(image_path)}")
        
        inputs = processor.apply_chat_template(
            messages, add_generation_prompt=True, tokenize=True,
            return_dict=True, return_tensors="pt"
        ).to(model.device, dtype=torch.bfloat16)
        
        input_len = inputs["input_ids"].shape[-1]
        
        with torch.inference_mode():
            generation = model.generate(**inputs, max_new_tokens=512, do_sample=False)
            generation = generation[0][input_len:]
        
        result = processor.decode(generation, skip_special_tokens=True)
        logger.info(f"VLM processing completed for {os.path.basename(image_path)}")
        return result
    except Exception as e:
        logger.error(f"Error processing image {image_path} with VLM: {e}")
        return None

def classify_image_theme(image_path, model, processor, device):
    """Classify image into one of the predefined themes"""
    try:
        messages = get_theme_classification_messages(image_path)
        theme_output = run_vlm_on_image(image_path, messages, model, processor, device)
        
        if theme_output:
            # Clean the output and check if it's in our themes list
            theme_output = theme_output.strip().lower()
            
            # Find the best matching theme
            for theme in image_themes:
                if theme.lower() in theme_output:
                    logger.info(f"Classified image as theme: {theme}")
                    return theme
            
            # If no exact match, try to find partial matches
            for theme in image_themes:
                if any(word in theme_output for word in theme.split()):
                    logger.info(f"Classified image as theme (partial match): {theme}")
                    return theme
            
            # If still no match, return the first theme as fallback
            logger.warning(f"Could not classify theme, using default. Output was: {theme_output}")
            return image_themes[0]
        else:
            logger.warning("Theme classification failed, using default theme")
            return image_themes[0]
    except Exception as e:
        logger.error(f"Error classifying image theme: {e}")
        return image_themes[0]

def parse_vlm_output(output):
    if not output:
        return []
    
    # Parse output into (question, count, property, objects) tuples
    lines = output.strip().split('\n')
    questions = []
    
    for line in lines:
        line = line.strip()
        if not line or not line[0].isdigit():
            continue
            
        try:
            # Expected format: 1. How many objects made of wood are present? 2 (physical) (bowls)
            # Split on '?' first
            if '?' not in line:
                continue
                
            q_part, rest = line.split('?', 1)
            
            # Extract question (remove number and dot)
            if '.' in q_part:
                question = q_part.split('.', 1)[1].strip() + '?'
            else:
                continue
                
            # Parse the rest: should be "Number (property) (objects)"
            rest = rest.strip()
            
            # Find parentheses
            parts = []
            current = ""
            in_parens = False
            paren_content = ""
            
            for char in rest:
                if char == '(':
                    if current.strip():
                        parts.append(current.strip())
                    current = ""
                    in_parens = True
                    paren_content = ""
                elif char == ')':
                    if in_parens:
                        parts.append(paren_content.strip())
                        paren_content = ""
                        in_parens = False
                elif in_parens:
                    paren_content += char
                else:
                    current += char
            
            if current.strip():
                parts.append(current.strip())
            
            # Extract count, property, and objects
            count = parts[0] if len(parts) > 0 else ""
            property_dim = parts[1] if len(parts) > 1 else ""
            objects = parts[2] if len(parts) > 2 else ""
            
            # Clean up count (remove any non-digit characters except the number)
            count = ''.join(filter(str.isdigit, count.split()[0])) if count else ""
            
            if question and count and property_dim:
                questions.append((question, count, property_dim, objects))
                logger.info(f"Parsed question: {question} -> {count} ({property_dim}) ({objects})")
            else:
                logger.warning(f"Incomplete parsing for line: {line}")
                
        except Exception as e:
            logger.warning(f"Failed to parse line: {line}. Error: {e}")
            continue
    
    return questions

def add_image_to_doc(doc, image_path):
    """Add image to document with error handling"""
    try:
        # Validate image first
        if not validate_image(image_path):
            logger.error(f"Skipping invalid image: {image_path}")
            doc.add_paragraph(f"[ERROR: Could not load image {os.path.basename(image_path)}]")
            return False
        
        doc.add_picture(image_path, width=Inches(3))
        logger.info(f"Added image to document: {os.path.basename(image_path)}")
        return True
    except Exception as e:
        logger.error(f"Error adding image {image_path} to document: {e}")
        doc.add_paragraph(f"[ERROR: Could not add image {os.path.basename(image_path)} - {str(e)}]")
        return False
    
def extract_image_number(filename):
    match = re.search(r'(\d+)', filename)
    return match.group(1) if match else filename

def process_type(type_name, image_dir, output_docx, model, processor, device, test_mode=False, max_images=None):
    """Process images with optional test mode"""
    logger.info(f"Starting processing for type: {type_name}")
    
    doc = Document()
    image_files = [f for f in os.listdir(image_dir) if f.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'))]
    
    # Natural sorting to handle image1, image2, ..., image10, image11 correctly
    def natural_sort_key(filename):
        import re
        # Extract numbers and convert to integers for proper sorting
        parts = re.split(r'(\d+)', filename)
        return [int(part) if part.isdigit() else part for part in parts]
    
    image_files = sorted(image_files, key=natural_sort_key)
    
    if test_mode and max_images:
        image_files = image_files[:max_images]
        logger.info(f"Test mode: Processing only first {max_images} images")
    
    logger.info(f"Found {len(image_files)} images to process")
    
    processed_count = 0
    error_count = 0
    
    for idx, img_file in enumerate(image_files, 1):
        img_path = os.path.join(image_dir, img_file)
        logger.info(f"Processing image {idx}/{len(image_files)}: {img_file}")
        
        # Classify image theme
        image_theme = classify_image_theme(img_path, model, processor, device)
        
        # Add header with image number and theme
        image_number = extract_image_number(img_file)
        doc.add_heading(f'Image {image_number}\t\t({image_theme})', level=2)
        
        # Try to add image to document
        image_added = add_image_to_doc(doc, img_path)
        
        if image_added:
            try:
                messages = get_messages(idx, img_file, img_path)
                vlm_output = run_vlm_on_image(img_path, messages, model, processor, device)
                
                if vlm_output:
                    questions = parse_vlm_output(vlm_output)
                    
                    if questions:
                        logger.info(f"Generated {len(questions)} questions for {img_file}")
                        for i, (q, count, prop, objs) in enumerate(questions, 1):
                            # doc.add_paragraph(f"{i}. {q} {count} ({prop}) ({objs})")
                            para = doc.add_paragraph(
                                f"{i}. {q}     {count}        ({prop})      ({objs})"
                            )
                            # para =doc.add_paragraph(
                            #     f"{q}     {count}        ({prop})      ({objs})",
                            #     style='List Number'
                            # )
                            # Remove spacing after
                            para.paragraph_format.space_after = 0
                    else:
                        logger.warning(f"No questions generated for {img_file}")
                        doc.add_paragraph("[No questions generated]")
                    
                    # Add raw VLM output for debugging
                    if test_mode:
                        doc.add_paragraph("Raw VLM Output:")
                        doc.add_paragraph(vlm_output)
                    
                    processed_count += 1
                else:
                    error_count += 1
                    doc.add_paragraph("[ERROR: VLM processing failed]")
                    
            except Exception as e:
                logger.error(f"Error processing {img_file}: {e}")
                error_count += 1
                doc.add_paragraph(f"[ERROR: {str(e)}]")
        else:
            error_count += 1
        
        doc.add_page_break()
        
        # Save periodically in test mode
        if test_mode and idx % 5 == 0:
            temp_output = f"temp_{output_docx}"
            doc.save(temp_output)
            logger.info(f"Saved temporary file: {temp_output}")
    
    doc.save(output_docx)
    logger.info(f"Saved final document: {output_docx}")
    logger.info(f"Processing complete for {type_name}: {processed_count} successful, {error_count} errors")
    
    return processed_count, error_count

def main():
    # Configuration
    TEST_MODE = False  # Set to False for full processing
    MAX_IMAGES_PER_TYPE = 1  # Number of images to test per type
    
    base_dir = "../merged_data"
    type_dirs = {
        "Real": os.path.join(base_dir, "Real"),
        "Animated": os.path.join(base_dir, "Animated"),
        "AI_Generated": os.path.join(base_dir, "AI_Generated"),
    }
    
    model_path = "/var/scratch/ave303/models/gemma-3-27b-it"
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    
    logger.info(f"Starting Gemma3 processing - Test Mode: {TEST_MODE}")
    logger.info(f"Loading model from {model_path} on device {device}")
    
    try:
        model = Gemma3ForConditionalGeneration.from_pretrained(
            model_path, 
            trust_remote_code=True, 
            torch_dtype=torch.bfloat16, 
            device_map="auto", 
            low_cpu_mem_usage=True
        ).eval()
        processor = AutoProcessor.from_pretrained(model_path, trust_remote_code=True)
        logger.info("Model and processor loaded successfully")
    except Exception as e:
        logger.error(f"Error loading model: {e}")
        return
    
    total_processed = 0
    total_errors = 0
    
    for type_name, dir_path in type_dirs.items():
        if not os.path.exists(dir_path):
            logger.warning(f"Directory does not exist: {dir_path}")
            continue
            
        output_docx = f"{'test_' if TEST_MODE else ''}{type_name}_questions.docx"
        
        try:
            processed, errors = process_type(
                type_name, dir_path, output_docx, model, processor, device, 
                test_mode=TEST_MODE, max_images=MAX_IMAGES_PER_TYPE if TEST_MODE else None
            )
            total_processed += processed
            total_errors += errors
        except Exception as e:
            logger.error(f"Error processing type {type_name}: {e}")
            total_errors += 1
    
    logger.info(f"Processing complete. Total processed: {total_processed}, Total errors: {total_errors}")

if __name__ == "__main__":
    main()